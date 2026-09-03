import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = 'TH Sarabun New'
    if level == 1:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0D, 0x2C, 0x5E) # Deep Navy
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x4F, 0xA0) # Royal Blue
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate Dark
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
    return h

def add_body_paragraph(doc, text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return p

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def create_complete_report():
    doc = Document()
    
    # Page Margins (A4 Standard)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ═════════════════════════════════════════════════════════════════════════
    # ปกหน้า (COVER PAGE)
    # ═════════════════════════════════════════════════════════════════════════
    p_top_gap = doc.add_paragraph()
    p_top_gap.paragraph_format.space_before = Pt(36)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title_p.add_run("รายงานการวิเคราะห์และออกแบบระบบ\n(System Analysis and Database Design Report)\n")
    r1.font.name = 'TH Sarabun New'
    r1.font.size = Pt(24)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x0D, 0x2C, 0x5E)

    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = proj_p.add_run("โครงการพัฒนาระบบแผนผังดิจิทัลและระบบนำทางอัจฉริยะ\nมหาวิทยาลัยราชภัฏศรีสะเกษ\n(SSKRU Campus Map & Smart Navigation System)\n")
    r2.font.name = 'TH Sarabun New'
    r2.font.size = Pt(20)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x1A, 0x4F, 0xA0)

    p_mid_gap = doc.add_paragraph()
    p_mid_gap.paragraph_format.space_before = Pt(80)

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = org_p.add_run("คณะศิลปศาสตร์และวิทยาศาสตร์ / สาขาวิชาเทคโนโลยีสารสนเทศ\nมหาวิทยาลัยราชภัฏศรีสะเกษ\n(Sisaket Rajabhat University)\nปีการศึกษา 2567 - 2569")
    r3.font.name = 'TH Sarabun New'
    r3.font.size = Pt(16)
    r3.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # บทที่ 1: บทนำ (INTRODUCTION)
    # ═════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "บทที่ 1: บทนำ (Introduction)", level=1)

    add_heading_styled(doc, "1.1 ความเป็นมาและความสำคัญของปัญหา (Background & Significance)", level=2)
    add_body_paragraph(doc, 
        "มหาวิทยาลัยราชภัฏศรีสะเกษ เป็นสถาบันอุดมศึกษาเพื่อการพัฒนาท้องถิ่นที่มีพื้นที่วิทยาเขตขนาดใหญ่ มีอาคารเรียน อาคารปฏิบัติการ "
        "สำนักงาน ศูนย์กีฬา หอพักนักศึกษา และหน่วยงานสนับสนุนการศึกษาจำนวนมากกว่า 20 อาคาร ซึ่งในแต่ละปีมีนักศึกษาใหม่ บุคลากรทางการศึกษา "
        "ตลอดจนประชาชนและผู้มาติดต่อราชการเป็นจำนวนมาก ผู้ใช้งานส่วนใหญ่มักประสบปัญหาไม่ทราบพิกัดที่ตั้งของอาคาร ไม่ทราบเส้นทางที่สั้นที่สุด "
        "และไม่ทราบว่าหน่วยงานหรือห้องเรียนที่ต้องการติดต่อตั้งอยู่ที่อาคารใด ชั้นใด"
    )
    add_body_paragraph(doc,
        "จากปัญหาดังกล่าว จึงได้มีแนวคิดในการพัฒนาระบบแผนผังดิจิทัล มรภ.ศรีสะเกษ (SSKRU Campus Map & Navigation System) "
        "โดยนำเทคโนโลยีเว็บสมัยใหม่ร่วมกับแผนที่สามมิติ (3D Interactive Map) และระบบดาวเทียมระบุตำแหน่งแบบเรียลไทม์ (Live GPS Geolocation) "
        "มาประยุกต์ใช้เพื่ออำนวยความสะดวกในการสืบค้นข้อมูลอาคาร การคำนวณเส้นทาง และการนำทางแบบเลี้ยวต่อเลี้ยว (Turn-by-Turn) สู่จุดหมายปลายทางอย่างแม่นยำ"
    )

    add_heading_styled(doc, "1.2 วัตถุประสงค์ของโครงการ (Objectives)", level=2)
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.left_indent = Inches(0.25)
    p_obj.paragraph_format.line_spacing = 1.15
    for item in [
        "1. เพื่อออกแบบและพัฒนาระบบแผนผังมหาวิทยาลัยแบบ 3D Interactive Map ที่รองรับการใช้งานบนคอมพิวเตอร์และสมาร์ตโฟน (Responsive Web Design)",
        "2. เพื่อพัฒนาระบบตรวจจับพิกัดดาวเทียม GPS สดของผู้ใช้งาน และแปลงพิกัดเข้าสู่แผนผังจำลองด้วยสมการการปรับเทียบเชิงเส้น (Coordinate Calibration)",
        "3. เพื่อเชื่อมโยงระบบการนำทางอัจฉริยะ (Smart Navigation) สู่ Google Maps นำทางผู้ใช้งานไปยังอาคารเป้าหมายได้อย่างถูกต้อง",
        "4. เพื่อสร้างระบบฐานข้อมูลกลางและระบบบริหารจัดการสำหรับผู้ดูแลระบบ (Admin Panel) ในการจัดการข้อมูลอาคาร ปักหมุด และสืบค้นสถิติการใช้งาน",
        "5. เพื่อพัฒนาระบบลงทะเบียนและยืนยันตัวตนที่สอดคล้องกับพระราชบัญญัติคุ้มครองข้อมูลส่วนบุคคล พ.ศ. 2562 (PDPA Compliance)"
    ]:
        run = p_obj.add_run(item + "\n")
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(16)

    add_heading_styled(doc, "1.3 ขอบเขตของโครงการ (Project Scope)", level=2)
    p_sc = doc.add_paragraph()
    p_sc.paragraph_format.left_indent = Inches(0.25)
    p_sc.paragraph_format.line_spacing = 1.15
    for item in [
        "• ขอบเขตด้านผู้ใช้งานทั่วไป (General Public / Guest): สามารถเปิดดูแผนผัง 3D, ค้นหาอาคาร, ดูรายละเอียดและรูปภาพ, กดจับพิกัด GPS ของตนเอง, และกดนำทางสู่ Google Maps ได้",
        "• ขอบเขตด้านนักศึกษา (Students): สามารถเข้าสู่ระบบด้วยรหัสนักศึกษา, ตรวจสอบข้อมูลสิทธิประโยชน์, บันทึกสถานที่โปรด และรีเซ็ตรหัสผ่านผ่าน OTP ทางอีเมลได้",
        "• ขอบเขตด้านอาจารย์และบุคลากร (Staff): สามารถลงทะเบียนด้วยอีเมลมหาวิทยาลัย (@sskru.ac.th), เข้าถึงข้อมูลห้องพักอาจารย์และเบอร์ติดต่อภายในได้",
        "• ขอบเขตด้านผู้ดูแลระบบ (System Admin): สามารถเข้าสู่ระบบหลังบ้านเพื่อเพิ่ม/ลบ/แก้ไขข้อมูลอาคาร, ลากหมุดเปลี่ยนตำแหน่งบนแผนผัง (Drag-and-Drop), อนุมัติสิทธิ์บุคลากร และดูรายงานสถิติผู้เข้าชม (Visitor Analytics)"
    ]:
        run = p_sc.add_run(item + "\n")
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(16)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # บทที่ 2: การวิเคราะห์และออกแบบฐานข้อมูล (DATABASE DESIGN)
    # ═════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "บทที่ 2: การวิเคราะห์และออกแบบฐานข้อมูล (Database Design)", level=1)

    add_heading_styled(doc, "2.1 แผนภาพความสัมพันธ์ของข้อมูล (Entity-Relationship Diagram : ER-Diagram)", level=2)
    add_body_paragraph(doc,
        "การออกแบบฐานข้อมูลของระบบ SSKRU Campus Map ได้รับการจัดทำตามหลักการกำหนดความสัมพันธ์เชิงสัมพันธ์ (Relational Database Design) "
        "โดยแบ่งโครงสร้างข้อมูลออกเป็น 8 Entity หลัก เพื่อรองรับการจัดเก็บข้อมูลอาคาร พิกัดแผนผัง บัญชีผู้ใช้งาน สถิติการเข้าชม "
        "และกิจกรรมการใช้งานบนแผนที่ ดังแสดงในแผนภาพ ER-Diagram รูปที่ 2.1:"
    )

    # Embed High-Res ER Diagram Image
    img_path = os.path.join(BASE_DIR, "images", "er_diagram_hd.png")
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6.4))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("รูปที่ 2.1: แผนภาพแสดงความสัมพันธ์ของข้อมูล (Entity-Relationship Diagram)")
        r_cap.font.name = 'TH Sarabun New'
        r_cap.font.size = Pt(14)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    add_heading_styled(doc, "2.2 พจนานุกรมข้อมูลและรายละเอียด Entity (Data Dictionary)", level=2)
    add_body_paragraph(doc,
        "พจนานุกรมข้อมูล (Data Dictionary) อธิบายโครงสร้างของทุกแอตทริบิวต์ (Attributes), ชนิดข้อมูล (Data Types), "
        "การกำหนดคีย์หลัก (Primary Key : PK), คีย์นอก (Foreign Key : FK), คีย์เฉพาะ (Unique Key : UK) และคำอธิบายการทำงานในระบบ ดังนี้:"
    )

    # Comprehensive Tables Definition
    tables = [
        ("2.2.1 Entity: Building (ข้อมูลอาคารและสถานที่)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับประจำแถว (Primary Key Auto-increment)"),
            ("building_id", "Integer", "UK", "No", "หมายเลขประจำอาคาร (เช่น 1, 2, 3... อ้างอิงตามแผนผัง)"),
            ("name", "Varchar(255)", "-", "No", "ชื่อทางการของอาคาร (ภาษาไทย)"),
            ("name_en", "Varchar(255)", "-", "Yes", "ชื่ออาคาร (ภาษาอังกฤษ)"),
            ("category", "Varchar(50)", "-", "No", "หมวดหมู่อาคาร (academic, office, facility, library)"),
            ("code", "Varchar(50)", "-", "Yes", "รหัสย่ออาคารสำหรับแสดงบนหมุด"),
            ("coord_x", "Float", "-", "No", "พิกัดแกน X บนรูปภาพแผนผัง 3D (Leaflet.js)"),
            ("coord_y", "Float", "-", "No", "พิกัดแกน Y บนรูปภาพแผนผัง 3D (Leaflet.js)"),
            ("lat", "Float", "-", "Yes", "พิกัดทางภูมิศาสตร์ ละติจูดดาวเทียม (GPS Latitude)"),
            ("lng", "Float", "-", "Yes", "พิกัดทางภูมิศาสตร์ ลองจิจูดดาวเทียม (GPS Longitude)"),
            ("description", "Text", "-", "Yes", "รายละเอียด ประวัติ และหน่วยงานภายในอาคาร"),
            ("phone", "Varchar(100)", "-", "Yes", "เบอร์โทรศัพท์ติดต่อภายในอาคาร"),
            ("tags", "JSON / Text", "-", "Yes", "คำสำคัญสำหรับการค้นหา (Keywords / Tags)"),
            ("image", "Varchar(500)", "-", "Yes", "URL หรือเส้นทางไฟล์รูปภาพจำลองอาคาร"),
            ("created_at", "DateTime", "-", "No", "วันเวลาที่เพิ่มข้อมูลอาคาร"),
            ("updated_at", "DateTime", "-", "No", "วันเวลาที่แก้ไขข้อมูลอาคารล่าสุด")
        ]),
        ("2.2.2 Entity: Student (ข้อมูลบัญชีนักศึกษา)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับบัญชีนักศึกษา (Primary Key)"),
            ("student_id", "Varchar(20)", "UK", "No", "รหัสประจำตัวนักศึกษา (เช่น 65123456789)"),
            ("name", "Varchar(255)", "-", "No", "ชื่อและนามสกุลของนักศึกษา"),
            ("year_level", "Integer", "-", "No", "ชั้นปีการศึกษา (1, 2, 3, 4...)"),
            ("password_hash", "Varchar(128)", "-", "Yes", "รหัสผ่านที่ผ่านการเข้ารหัสลับแบบ SHA-256"),
            ("password_plain", "Varchar(128)", "-", "Yes", "รหัสผ่านสำหรับระบบจัดการภายใน"),
            ("is_active", "Boolean", "-", "No", "สถานะการเปิดใช้งานบัญชี (True = ใช้งานได้)"),
            ("created_at", "DateTime", "-", "No", "วันเวลาที่ลงทะเบียนเข้าใช้งาน")
        ]),
        ("2.2.3 Entity: StaffUser (ข้อมูลอาจารย์และบุคลากร)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับบัญชีบุคลากร (Primary Key)"),
            ("username", "Varchar(100)", "UK", "No", "ชื่อผู้ใช้งานสำหรับเข้าสู่ระบบ (Username)"),
            ("email", "Varchar(255)", "UK", "No", "อีเมลมหาวิทยาลัยของบุคลากร (@sskru.ac.th)"),
            ("password_hash", "Varchar(128)", "-", "No", "รหัสผ่านเข้ารหัสลับแบบ SHA-256"),
            ("password_plain", "Varchar(128)", "-", "Yes", "รหัสผ่าน"),
            ("is_active", "Boolean", "-", "No", "สถานะเปิดใช้งานบัญชี"),
            ("is_approved", "Boolean", "-", "No", "สถานะการตรวจสอบและอนุมัติสิทธิ์โดยแอดมิน"),
            ("created_at", "DateTime", "-", "No", "วันเวลาที่สมัครสมาชิก")
        ]),
        ("2.2.4 Entity: VisitorLog (ข้อมูลผู้เข้าชมเว็บไซต์)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับประวัติการเข้าชม (Primary Key)"),
            ("ip_address", "GenericIP", "-", "Yes", "หมายเลข IP Address ของผู้เข้าชม"),
            ("user_agent", "Text", "-", "Yes", "ข้อมูล Web Browser และระบบอุปกรณ์ของผู้เข้าชม"),
            ("device_type", "Varchar(20)", "-", "No", "ประเภทอุปกรณ์ที่ใช้งาน (mobile, desktop, tablet)"),
            ("os_name", "Varchar(50)", "-", "Yes", "ชื่อระบบปฏิบัติการ (iOS, Android, macOS, Windows)"),
            ("browser", "Varchar(50)", "-", "Yes", "ชื่อโปรแกรมเบราว์เซอร์ (Safari, Chrome, Edge)"),
            ("page_path", "Varchar(255)", "-", "No", "เส้นทางหน้าเว็บที่เปิดเข้าชม (เช่น /, /login/)"),
            ("referrer", "Varchar(500)", "-", "Yes", "URL ต้นทางที่ส่งผู้ใช้งานมา"),
            ("session_id", "Varchar(64)", "-", "Yes", "รหัสเซสชันเฉพาะตัวเพื่อแยกผู้เยี่ยมชม"),
            ("timestamp", "DateTime", "-", "No", "วันเวลาที่เข้าชมระบบ")
        ]),
        ("2.2.5 Entity: UserEvent (กิจกรรมการใช้งานบนแผนที่)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับกิจกรรม (Primary Key)"),
            ("visitor_id", "Integer", "FK", "Yes", "รหัสอ้างอิงไปยัง VisitorLog.id (Foreign Key)"),
            ("event_type", "Varchar(30)", "-", "No", "ประเภทกิจกรรม (page_view, search, navigate, share)"),
            ("event_data", "Varchar(255)", "-", "Yes", "รายละเอียดกิจกรรม เช่น ชื่ออาคารที่กดนำทาง หรือคำค้นหา"),
            ("timestamp", "DateTime", "-", "No", "วันเวลาที่เกิดกิจกรรมการใช้งาน")
        ]),
        ("2.2.6 Entity: UserActivityLog (ประวัติกิจกรรมและความปลอดภัย)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับบันทึกกิจกรรม (Primary Key)"),
            ("user_id", "Varchar(100)", "-", "No", "รหัสประจำตัวผู้ใช้งาน (รหัสนักศึกษา หรือ Username)"),
            ("user_name", "Varchar(255)", "-", "No", "ชื่อและนามสกุลของผู้ใช้งาน"),
            ("role", "Varchar(50)", "-", "No", "บทบาทสิทธิ์ในระบบ (student, staff, admin)"),
            ("email", "Varchar(255)", "-", "Yes", "อีเมลของผู้ใช้งาน"),
            ("ip_address", "Varchar(100)", "-", "Yes", "หมายเลข IP Address ขณะล็อกอิน"),
            ("device", "Varchar(255)", "-", "Yes", "ข้อมูลอุปกรณ์และระบบปฏิบัติการ"),
            ("timestamp", "DateTime", "-", "No", "วันเวลาที่ทำการเข้าสู่ระบบหรือทำกิจกรรม")
        ]),
        ("2.2.7 Entity: PasswordResetToken (โทเค็นและ OTP กู้คืนรหัสผ่าน)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับโทเค็น (Primary Key)"),
            ("user_type", "Varchar(20)", "-", "No", "ประเภทผู้ใช้ที่ขอกู้คืนรหัส (student หรือ staff)"),
            ("identifier", "Varchar(100)", "-", "No", "รหัสนักศึกษา หรือ Username บุคลากร"),
            ("email", "Varchar(255)", "-", "No", "อีเมลที่ใช้รับรหัส OTP สำหรับยืนยัน"),
            ("token", "Varchar(64)", "UK", "No", "รหัสโทเค็นความปลอดภัยแบบสุ่ม (Unique Token)"),
            ("otp", "Varchar(10)", "-", "No", "รหัสผ่านชั่วคราว OTP 6 หลัก"),
            ("expires_at", "DateTime", "-", "No", "วันเวลาหมดอายุของรหัส OTP"),
            ("used", "Boolean", "-", "No", "สถานะการนำโทเค็นไปใช้งาน (True = ใช้งานแล้ว)"),
            ("created_at", "DateTime", "-", "No", "วันเวลาที่สร้างคำขอกู้คืนรหัสผ่าน")
        ]),
        ("2.2.8 Entity: AdminSession (เซสชันความปลอดภัยผู้ดูแลระบบ)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบายความหมาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับเซสชันแอดมิน (Primary Key)"),
            ("token", "Varchar(128)", "UK", "No", "ค่าแฮชโทเค็นความปลอดภัยสำหรับเข้าถึง Admin Panel"),
            ("created_at", "DateTime", "-", "No", "วันเวลาที่สร้างเซสชัน"),
            ("expires_at", "DateTime", "-", "No", "วันเวลาหมดอายุของเซสชัน"),
            ("is_active", "Boolean", "-", "No", "สถานะการเปิดใช้งานเซสชัน")
        ])
    ]

    for title, rows in tables:
        add_heading_styled(doc, title, level=3)
        table = doc.add_table(rows=len(rows), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        # Header Row Styling
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(rows[0]):
            hdr_cells[idx].text = text
            set_cell_background(hdr_cells[idx], "1A4FA0") # Deep Blue
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'TH Sarabun New'
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(14)
            set_cell_margins(hdr_cells[idx], 100, 100, 120, 120)

        # Data Rows Styling
        for r_idx, row_data in enumerate(rows[1:], start=1):
            row_cells = table.rows[r_idx].cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = val
                set_cell_background(row_cells[c_idx], bg_color)
                p = row_cells[c_idx].paragraphs[0]
                if c_idx in [0, 2, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'TH Sarabun New'
                    run.font.size = Pt(13)
                    if c_idx == 2 and val in ["PK", "FK", "UK"]:
                        run.font.bold = True
                        if val == "PK":
                            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26) # Red
                        elif val == "FK":
                            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Blue
                        else:
                            run.font.color.rgb = RGBColor(0x05, 0x96, 0x69) # Green
                set_cell_margins(row_cells[c_idx], 80, 80, 100, 100)

        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # บทที่ 3: สถาปัตยกรรมระบบและการทำงาน (SYSTEM ARCHITECTURE)
    # ═════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "บทที่ 3: สถาปัตยกรรมระบบและการทำงาน (System Architecture)", level=1)

    add_heading_styled(doc, "3.1 สถาปัตยกรรมระดับซอฟต์แวร์ (Software Architecture Stack)", level=2)
    p_st = doc.add_paragraph()
    p_st.paragraph_format.left_indent = Inches(0.25)
    p_st.paragraph_format.line_spacing = 1.15
    for item in [
        "1. Presentation Layer (Frontend): พัฒนาด้วยเทคโนโลยี HTML5, Vanilla CSS3 พร้อมระบบ Design Tokens และ Glassmorphism UI, JavaScript ES6+ และใช้ Leaflet.js เป็นเอนจินหลักในการเรนเดอร์แผนผัง 3D พิกัดพิกเซล",
        "2. Application Logic Layer (Backend): พัฒนาด้วย Django Framework (Python 3.13) และสถาปัตยกรรม PHP API สลับการทำงานได้ทั้งสองระบบ รองรับการประมวลผล RESTful JSON Endpoints",
        "3. Data Persistence Layer (Database): รองรับทั้งระบบ Relational Database (PostgreSQL / SQLite) และ JSON File Datastore เป็นระบบสำรอง (Fallback Strategy)",
        "4. Deployment & Cloud Layer: รองรับการ Deploy บน Vercel Serverless Architecture พร้อมระบบการจัดการเซสชันแบบ Signed Cookies และระบบความปลอดภัย HTTPS Proxy SSL"
    ]:
        run = p_st.add_run(item + "\n")
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(16)

    add_heading_styled(doc, "3.2 อัลกอริทึมการแปลงพิกัดดาวเทียมสู่แผนผัง 3D (Coordinate Calibration)", level=2)
    add_body_paragraph(doc,
        "เนื่องจากแผนที่แบบจำลอง 3D ของมหาวิทยาลัยเป็นภาพกราฟิกแบบพิกเซล (Pixel Coordinate System) ขนาด 1536 x 1024 พิกเซล "
        "ซึ่งแตกต่างจากระบบพิกัดดาวเทียมโลกจริง (WGS84 GPS Latitude/Longitude) ระบบจึงได้ใช้อัลกอริทึมการปรับเทียบเชิงเส้นสองมิติ (2D Linear Interpolation) "
        "โดยอ้างอิงจากจุดควบคุมหลักภายในมหาวิทยาลัย (หอประชุม 80 พรรษา และ อาคารศูนย์ภาษาและคอมพิวเตอร์):"
    )
    p_math = doc.add_paragraph()
    p_math.paragraph_format.left_indent = Inches(0.5)
    p_math.paragraph_format.line_spacing = 1.15
    r_m = p_math.add_run(
        "KY = (Y2 - Y1) / (LAT2 - LAT1)\n"
        "KX = (X2 - X1) / (LNG2 - LNG1)\n\n"
        "Y_map = Y1 + KY * (Latitude_GPS - LAT1)\n"
        "X_map = X1 + KX * (Longitude_GPS - LNG1)"
    )
    r_m.font.name = 'Courier New'
    r_m.font.size = Pt(13)
    r_m.font.bold = True
    r_m.font.color.rgb = RGBColor(0x1A, 0x4F, 0xA0)

    add_heading_styled(doc, "3.3 ระบบการนำทางอัจฉริยะ (Smart Turn-by-Turn Navigation)", level=2)
    add_body_paragraph(doc,
        "เมื่อผู้ใช้งานเลือกอาคารปลายทางและกดปุ่ม 'นำทาง (Google Maps)' ระบบจะตรวจสอบพิกัด GPS สดของผู้ใช้งาน "
        "หากตรวจพบพิกัด จะสร้าง Deep Link ส่งพิกัดเริ่มต้น (Origin) และพิกัดอาคารเป้าหมาย (Destination) ไปยังแอปพลิเคชัน Google Maps "
        "โดยเปิดหน้าต่างนำทางแบบเลี้ยวต่อเลี้ยวทันทีโดยไม่ถูกเบราว์เซอร์บล็อกป๊อปอัป ช่วยให้ผู้ใช้งานสามารถเดินเท้า ขี่จักรยานยนต์ หรือขับรถยนต์ไปยังอาคารได้อย่างสะดวก"
    )

    add_heading_styled(doc, "3.4 การปฏิบัติตาม พ.ร.บ. คุ้มครองข้อมูลส่วนบุคคล (PDPA Compliance)", level=2)
    add_body_paragraph(doc,
        "ระบบได้รับการออกแบบให้สอดคล้องกับ พ.ร.บ. คุ้มครองข้อมูลส่วนบุคคล พ.ศ. 2562 (PDPA) อย่างเคร่งครัด "
        "โดยมีกลไกขอความยินยอม (Consent Modal) ในการเข้าถึงตำแหน่ง GPS ของอุปกรณ์ก่อนการใช้งาน และในหน้าสมัครสมาชิก "
        "กำหนดให้ผู้ใช้เป็นผู้พิจารณาเลือกติ๊กยินยอมด้วยตนเอง (Unchecked by default) พร้อมทั้งจัดเก็บประวัติการยินยอมใน Local Storage อย่างโปร่งใส"
    )

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # บทที่ 4: สรุปผลและการต่อยอด (CONCLUSION)
    # ═════════════════════════════════════════════════════════════════════════
    add_heading_styled(doc, "บทที่ 4: สรุปผลการดำเนินงานและแนวทางการต่อยอด", level=1)
    add_body_paragraph(doc,
        "การพัฒนาระบบแผนผังดิจิทัลและระบบนำทางอัจฉริยะ มหาวิทยาลัยราชภัฏศรีสะเกษ (SSKRU Campus Map) ได้บรรลุวัตถุประสงค์ที่กำหนดไว้ "
        "สามารถให้บริการค้นหาอาคาร แสดงรายละเอียดสถานที่ ปักหมุดระบุตำแหน่งผู้ใช้งานสด และเปิดระบบนำทางแบบเลี้ยวต่อเลี้ยวได้อย่างมีประสิทธิภาพ "
        "โดยมีระบบฐานข้อมูลที่มั่นคง ปลอดภัย และมีส่วนต่อประสานผู้ใช้ (User Interface) ที่ทันสมัย ตอบสนองรวดเร็วบนทุกอุปกรณ์"
    )
    add_heading_styled(doc, "4.1 แนวทางการพัฒนาต่อยอดในอนาคต (Future Enhancements)", level=2)
    p_fut = doc.add_paragraph()
    p_fut.paragraph_format.left_indent = Inches(0.25)
    p_fut.paragraph_format.line_spacing = 1.15
    for item in [
        "1. การพัฒนาระบบแผนที่นำทางภายในอาคารแบบแยกชั้น (Indoor Floor-by-Floor Navigation)",
        "2. การเชื่อมโยงระบบเข้ากับตารางเรียนตารางสอนและปฏิทินการศึกษาของนักศึกษาแบบเรียลไทม์",
        "3. การประยุกต์ใช้เทคโนโลยีภาพเสมือนจริง (Augmented Reality : AR Navigation) สำหรับกล้องสมาร์ตโฟนในการนำทางบนถนนจริงภายในวิทยาเขต"
    ]:
        run = p_fut.add_run(item + "\n")
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(16)

    # Save final docx
    output_path = os.path.join(BASE_DIR, "docs", "SSKRU_Campus_Map_Report.docx")
    doc.save(output_path)
    print(f"Complete report generated at: {output_path}")

if __name__ == "__main__":
    create_complete_report()
