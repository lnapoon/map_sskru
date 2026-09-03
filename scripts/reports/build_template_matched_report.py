import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_font_run(run, font_name="Angsana New", size_pt=16, bold=False, color_rgb=None):
    run.font.name = font_name
    # Also set eastAsia/cs font for Thai in Word
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_p(doc, text="", font_name="Angsana New", size_pt=16, bold=False, align=None, space_before=0, space_after=0, indent=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        set_font_run(run, font_name, size_pt, bold)
    return p

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

def create_template_matched_report():
    doc = Document()

    # Match Page Setup (Standard A4)
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # ═════════════════════════════════════════════════════════════════════════
    # 1. หน้าปก (COVER PAGE) — Exactly matching 6712732121.docx style
    # ═════════════════════════════════════════════════════════════════════════
    
    # University Logo
    logo_path = os.path.join(BASE_DIR, "images", "image1.png")
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(10)
        p_logo.paragraph_format.space_after = Pt(14)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.25))

    add_p(doc, "รายงาน", font_name="Angsana New", size_pt=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "เรื่อง ระบบแผนผังดิจิทัลและระบบนำทางอัจฉริยะ มหาวิทยาลัยราชภัฏศรีสะเกษ", font_name="Angsana New", size_pt=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_p(doc, "จัดทำโดย", font_name="Angsana New", size_pt=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "นายมนต์พระกาฬ ทิมเกิด รหัสนักศึกษา 6712732120", font_name="Angsana New", size_pt=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_p(doc, "เสนอ", font_name="Angsana New", size_pt=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "ผู้ช่วยศาสตราจารย์ ดร.กนิษฐา อินธิชิต", font_name="Angsana New", size_pt=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    add_p(doc, "รายงานนี้เป็นส่วนหนึ่งของรายวิชาระบบฐานข้อมูล ( 4123202 )", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "ภาคเรียนที่ 2 ปีการศึกษา 2568", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "สาขาวิทยาการคอมพิวเตอร์ คณะศิลปศาสตร์และวิทยาศาสตร์", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "มหาวิทยาลัยราชภัฏศรีสะเกษ", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 2. สารบัญ (TABLE OF CONTENTS) — Exactly matching 6712732121.docx style
    # ═════════════════════════════════════════════════════════════════════════
    add_p(doc, "สารบัญ", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    toc_items = [
        ("Requirement", "1"),
        ("แผนภาพความสัมพันธ์ของข้อมูล (E-R Diagram)", "2"),
        ("ข้อมูลตารางทั้งหมด (Data Dictionary)", "4"),
        ("User Interface (UI)", "6")
    ]
    for title, page_num in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.line_spacing = 1.2
        r_title = p_toc.add_run(title)
        set_font_run(r_title, font_name="Angsana New", size_pt=18, bold=False)
        
        # Dot leaders spacing
        dots_count = max(4, 90 - len(title) * 2)
        r_dots = p_toc.add_run(" " + "." * dots_count + " ")
        set_font_run(r_dots, font_name="Angsana New", size_pt=14, bold=False)

        r_page = p_toc.add_run(page_num)
        set_font_run(r_page, font_name="Angsana New", size_pt=18, bold=True)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 3. หน้า 1: REQUIREMENT — Based on SSKRU Map System
    # ═════════════════════════════════════════════════════════════════════════
    add_p(doc, "1", font_name="Angsana New", size_pt=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "Requirement", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    add_p(doc, "ระบบแผนผังดิจิทัลและระบบนำทางอัจฉริยะ มหาวิทยาลัยราชภัฏศรีสะเกษ (SSKRU Campus Map)", font_name="Angsana New", size_pt=16, bold=True, space_after=4)
    add_p(doc, "• สามารถสืบค้นและแสดงตำแหน่งอาคารสถานที่ภายในมหาวิทยาลัยบนแผนผัง 3D ได้", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• สามารถค้นหาอาคารได้อย่างสะดวกรวดเร็วผ่านระบบค้นหาอัจฉริยะและตัวกรองหมวดหมู่ (อาคารเรียน, สำนักงาน, สิ่งอำนวยความสะดวก, หอสมุด)", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• สามารถจับพิกัดดาวเทียม GPS สดของผู้ใช้งานจริง และแปลงพิกัดมาปักหมุดเรดาร์สีฟ้าบนแผนผังจำลอง 3D ได้อย่างแม่นยำ", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• สามารถเชื่อมต่อระบบนำทางแบบเลี้ยวต่อเลี้ยว (Turn-by-Turn) ส่งพิกัดไปยัง Google Maps เพื่อนำทางผู้ใช้ไปยังอาคารเป้าหมายได้ทันที", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• สามารถเลือกรูปแบบการเดินทางได้หลากหลาย เช่น เดินเท้า, จักรยานยนต์, รถยนต์", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• มีระบบเข้าสู่ระบบสำหรับนักศึกษา (Student Login) และระบบสมัครสมาชิกสำหรับอาจารย์/บุคลากร (Staff Register)", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• มีระบบกู้คืนรหัสผ่านความปลอดภัยสูงผ่านการยืนยันรหัส OTP 6 หลักทางอีเมลมหาวิทยาลัย", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• มีระบบคุ้มครองข้อมูลส่วนบุคคล (PDPA Consent) ขออนุญาตการเข้าถึงพิกัดและข้อมูลตามกฎหมาย", font_name="Angsana New", size_pt=16, bold=False)
    add_p(doc, "• มีระบบผู้ดูแลระบบ (Admin Panel) สำหรับเพิ่ม ลบ แก้ไขข้อมูลอาคาร และลากย้ายพิกัดหมุด (Drag & Drop) บนแผนที่ได้โดยตรง", font_name="Angsana New", size_pt=16, bold=False, space_after=8)

    add_p(doc, "การทำงานของระบบ", font_name="Angsana New", size_pt=16, bold=True, space_after=2)
    add_p(doc, "หน้าแรกเมื่อเปิดเข้ามาจะแสดงแผนผัง 3D Interactive Map ของมหาวิทยาลัยราชภัฏศรีสะเกษ พร้อมแถบการ์ดรายการอาคาร Carousel ด้านล่าง และมีแถบค้นหาด้านบน ผู้ใช้สามารถกดปุ่มเป้าเล็ง GPS เพื่อดูว่าตนเองยืนอยู่จุดใดของมหาวิทยาลัย และเมื่อคลิกเลือกอาคาร ระบบจะเปิดแผงข้อมูลแสดงรูปภาพ รายละเอียด สถานะเปิด-ปิด เบอร์โทรศัพท์ และปุ่มนำทางสู่ Google Maps", font_name="Angsana New", size_pt=16, indent=True)

    add_p(doc, "วิธีการใช้งาน", font_name="Angsana New", size_pt=16, bold=True, space_after=2)
    add_p(doc, "1. ค้นหาอาคารที่ต้องการผ่านแถบค้นหา หรือคลิกเลือกหมวดหมู่อาคาร\n2. คลิกที่หมุดอาคารหรือการ์ดอาคารเพื่อเปิดดูข้อมูลและเวลาทำการ\n3. กดปุ่มเป้าเล็ง GPS ด้านขวาบนเพื่อระบุพิกัดตำแหน่งปัจจุบันของคุณ\n4. กดปุ่ม 'นำทาง (Google Maps)' เพื่อเปิดแอปแผนที่นำทางไปยังอาคารนั้นทันที", font_name="Angsana New", size_pt=16, indent=True, space_after=6)

    add_p(doc, "ข้อมูลหลัก", font_name="Angsana New", size_pt=16, bold=True, space_after=2)
    add_p(doc, "• ผู้ใช้ทั่วไป: ค้นหาอาคาร ดูรายละเอียด ตรวจจับพิกัด GPS ของตนเอง และกดนำทางสู่ Google Maps\n• นักศึกษา: เข้าสู่ระบบเพื่อเข้าถึงข้อมูลสิทธิประโยชน์ และจัดการข้อมูลส่วนตัว\n• อาจารย์/บุคลากร: สมัครสมาชิกเพื่อเข้าถึงข้อมูลหน่วยงานภายใน และรอการอนุมัติสิทธิ์\n• ผู้ดูแลระบบ (Admin): เพิ่ม แก้ไข ลบข้อมูลอาคาร ลากย้ายตำแหน่งหมุด อนุมัติสิทธิ์บุคลากร และดูสถิติผู้เข้าชม", font_name="Angsana New", size_pt=16, indent=True, space_after=6)

    add_p(doc, "ฟังก์ชันหลัก", font_name="Angsana New", size_pt=16, bold=True, space_after=2)
    add_p(doc, "• ระบบแผนผัง 3D Interactive — รองรับการซูม เลื่อน ย้ายมุมมองบนภาพความละเอียดสูง\n• ระบบระบุตำแหน่ง Live GPS — ตรวจจับพิกัดสดและปรับเทียบเข้าสู่แผนผังจำลอง\n• ระบบนำทาง Turn-by-Turn — ส่งต่อพิกัดอย่างแม่นยำสู่ Google Maps ภายนอก\n• ระบบจัดการฐานข้อมูลอาคาร — อัปเดตข้อมูลและพิกัดแบบเรียลไทม์\n• มีฟังก์ชันรองรับการใช้งานบนมือถือ (Responsive Web Design บน iOS และ Android)\n• รูปแบบเว็บทันสมัย สวยงาม ใช้งานง่าย โทนสีสบายตา และรองรับ Dark Mode / Light Mode", font_name="Angsana New", size_pt=16, indent=True)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 4. หน้า 2 - 3: แผนภาพความสัมพันธ์ของข้อมูล (E-R DIAGRAM)
    # ═════════════════════════════════════════════════════════════════════════
    add_p(doc, "2", font_name="Angsana New", size_pt=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "แผนภาพความสัมพันธ์ของข้อมูล (E-R Diagram)", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Embed ER Diagram image
    erd_img = os.path.join(BASE_DIR, "images", "er_diagram_hd.png")
    if os.path.exists(erd_img):
        p_erd = doc.add_paragraph()
        p_erd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_erd.paragraph_format.space_before = Pt(4)
        p_erd.paragraph_format.space_after = Pt(8)
        run_erd = p_erd.add_run()
        run_erd.add_picture(erd_img, width=Inches(6.4))

    add_p(doc, "1. เอนทิตี (Entity Description)", font_name="Angsana New", size_pt=16, bold=True, space_after=4)
    add_p(doc, "ระบบประกอบด้วย 8 เอนทิตีหลักที่ทำงานร่วมกัน ดังนี้:", font_name="Angsana New", size_pt=16)
    add_p(doc, "• buildings: เก็บข้อมูลอาคารและสถานที่ทั้งหมด พิกัดแผนที่ 3D (X, Y) และพิกัดดาวเทียม GPS จริง (Latitude, Longitude)", font_name="Angsana New", size_pt=16)
    add_p(doc, "• students: เก็บข้อมูลบัญชีนักศึกษา รหัสนักศึกษา ชื่อ-นามสกุล ชั้นปี และรหัสผ่าน", font_name="Angsana New", size_pt=16)
    add_p(doc, "• staff_users: เก็บข้อมูลบัญชีอาจารย์และบุคลากร ชื่อผู้ใช้ อีเมล รหัสผ่าน และสถานะการอนุมัติสิทธิ์", font_name="Angsana New", size_pt=16)
    add_p(doc, "• visitor_logs: เก็บข้อมูลสถิติผู้เข้าชมเว็บไซต์ หมายเลข IP ข้อมูลอุปกรณ์ ระบบปฏิบัติการ เบราว์เซอร์ และรหัสเซสชัน", font_name="Angsana New", size_pt=16)
    add_p(doc, "• user_events: บันทึกกิจกรรมการใช้งานบนแผนที่ เช่น การกดค้นหา การกดเลือกดูอาคาร และการกดนำทาง GPS", font_name="Angsana New", size_pt=16)
    add_p(doc, "• user_activity_logs: บันทึกประวัติการเข้าสู่ระบบและกิจกรรมความปลอดภัยของผู้ใช้งานทุกสิทธิ์", font_name="Angsana New", size_pt=16)
    add_p(doc, "• password_reset_tokens: จัดเก็บรหัส OTP 6 หลัก และโทเค็นสำหรับการกู้คืนรหัสผ่านผ่านอีเมล", font_name="Angsana New", size_pt=16)
    add_p(doc, "• admin_sessions: จัดเก็บโทเค็นความปลอดภัยสำหรับเซสชันการเข้าใช้งานระบบของผู้ดูแลระบบ (Admin)", font_name="Angsana New", size_pt=16, space_after=8)

    doc.add_page_break()

    # หน้า 3: คำอธิบายความสัมพันธ์
    add_p(doc, "3", font_name="Angsana New", size_pt=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "2. คำอธิบายความสัมพันธ์ (Relationship Analysis)", font_name="Angsana New", size_pt=16, bold=True, space_after=6)
    add_p(doc, "ความสัมพันธ์ระหว่างข้อมูลถูกออกแบบตามหลักการจัดการฐานข้อมูลเชิงสัมพันธ์ (Normalization) ดังนี้:", font_name="Angsana New", size_pt=16, space_after=6)

    relationships = [
        ("visitor_logs ─── user_events (TRACKS): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "ผู้เข้าชม 1 เซสชัน (1 Visitor) สามารถทำกิจกรรมบนแผนที่ได้หลายกิจกรรม (Many Events) เช่น ค้นหาอาคาร 2 ครั้ง และกดนำทาง 1 ครั้ง โดยใช้ visitor_id เป็น Foreign Key เชื่อมโยง"),
        ("students ─── password_reset_tokens (REQUESTS): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "นักศึกษา 1 คน สามารถส่งคำขอกู้คืนรหัสผ่านเพื่อรับรหัส OTP ได้หลายครั้งตามช่วงเวลาที่ลืมรหัส โดยใช้ identifier เป็นตัวอ้างอิงรหัสนักศึกษา"),
        ("staff_users ─── password_reset_tokens (REQUESTS): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "บุคลากร 1 ท่าน สามารถขอรับรหัส OTP เพื่อรีเซ็ตรหัสผ่านได้หลายครั้ง โดยใช้ identifier อ้างอิง Username หรือ Email"),
        ("students ─── user_activity_logs (LOGS): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "นักศึกษา 1 คน มีประวัติการเข้าสู่ระบบและทำกิจกรรมในระบบได้หลายครั้ง โดยบันทึกรหัสนักศึกษาลงใน user_id"),
        ("staff_users ─── user_activity_logs (LOGS): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "บุคลากร 1 ท่าน มีประวัติการล็อกอินและการใช้งานระบบบันทึกไว้หลายรายการ โดยเชื่อมโยงผ่าน user_id"),
        ("staff_users ─── admin_sessions (AUTHENTICATES): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "บุคลากรที่มีสิทธิ์ผู้ดูแลระบบ สามารถสร้างเซสชันความปลอดภัย Admin Token ได้เมื่อเข้าสู่ระบบหลังบ้าน"),
        ("buildings ─── user_events (INTERACTS): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "อาคาร 1 อาคาร สามารถถูกผู้ใช้งานหลายคนกดค้นหา คลิกดูข้อมูล หรือกดนำทางได้หลายครั้ง โดยบันทึกรหัสอาคารลงใน event_data"),
        ("buildings ─── user_activity_logs (UPDATES): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "อาคาร 1 อาคาร สามารถมีประวัติการแก้ไขข้อมูล หรือลากย้ายพิกัดโดยผู้ดูแลระบบได้หลายครั้ง"),
        ("admin_sessions ─── user_activity_logs (AUDITS): ความสัมพันธ์แบบ One-to-Many (1:N)",
         "เซสชันแอดมิน 1 เซสชัน สามารถมีบันทึก Audit Log กิจกรรมการจัดการระบบได้หลายรายการ")
    ]

    for rel_title, rel_desc in relationships:
        add_p(doc, rel_title, font_name="Angsana New", size_pt=16, bold=True)
        add_p(doc, rel_desc, font_name="Angsana New", size_pt=16, indent=True, space_after=4)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 5. หน้า 4 - 5: ข้อมูลตารางทั้งหมด (DATA DICTIONARY) — Exact 6712732121 format
    # ═════════════════════════════════════════════════════════════════════════
    add_p(doc, "4", font_name="Angsana New", size_pt=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "ข้อมูลตารางทั้งหมด (Data Dictionary)", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_p(doc, "Entity ทั้งหมด", font_name="Angsana New", size_pt=16, bold=True, space_after=4)

    # Table 1: buildings
    add_p(doc, "ตารางข้อมูลอาคารและสถานที่ (buildings)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ใช้เก็บข้อมูลอาคารสถานที่ทั้งหมดภายในมหาวิทยาลัย เพื่อใช้แสดงผลบนแผนผัง 3D และนำทาง GPS", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: ประกอบด้วยรหัสประจำตัว (id: PK), หมายเลขอาคาร (building_id: UK), ชื่ออาคารไทย (name), ชื่ออังกฤษ (name_en), หมวดหมู่ (category), รหัสย่อ (code), พิกัดแผนที่ (coord_x, coord_y), พิกัดดาวเทียมจริง (lat, lng), รายละเอียด (description), เบอร์โทร (phone) และรูปภาพ (image)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: มีรายการอาคาร 20 อาคาร เช่น อาคาร 1 อาคารเฉลิมพระเกียรติ 80 พรรษา (หมวดหมู่ Academic, พิกัด [637, 644], GPS [15.119754, 104.358862]), อาคาร 9 หอสมุด (หมวดหมู่ Library), อาคาร 14 สำนักงานอธิการบดี (หมวดหมู่ Office) เป็นต้น", font_name="Angsana New", size_pt=16, space_after=6)

    # Table 2: students_roster / students
    add_p(doc, "ตารางข้อมูลนักศึกษา (students_roster / students)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ใช้จัดการข้อมูลบัญชีนักศึกษาสำหรับการตรวจสอบสิทธิ์เข้าใช้งานระบบแผนที่", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: ประกอบด้วยรหัสแถว (id: PK), รหัสนักศึกษา (student_id: UK), ชื่อ-นามสกุล (name), ชั้นปี (year_level), รหัสผ่านเข้ารหัส (password_hash) และสถานะเปิดใช้งาน (is_active)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: เช่น รหัสนักศึกษา 65123456789 ชื่อ นายสมชาย ใจดี ชั้นปีที่ 2 สถานะ Active พร้อมใช้งาน", font_name="Angsana New", size_pt=16, space_after=6)

    # Table 3: users / staff_users
    add_p(doc, "ตารางข้อมูลบุคลากรและอาจารย์ (users / staff_users)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ใช้เก็บข้อมูลบัญชีอาจารย์และบุคลากรสำหรับเข้าใช้งานระบบและรับสิทธิ์ดูแลข้อมูล", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: ประกอบด้วยรหัสบัญชี (id: PK), ชื่อผู้ใช้ (username: UK), อีเมลมหาวิทยาลัย (email: UK), รหัสผ่านเข้ารหัส (password_hash), สถานะใช้งาน (is_active) และสถานะอนุมัติสิทธิ์โดยแอดมิน (is_approved)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: เช่น Username: ajarn_somchai, Email: somchai@sskru.ac.th, สถานะ Approved (อนุมัติแล้ว)", font_name="Angsana New", size_pt=16, space_after=6)

    # Table 4: visitor_logs
    add_p(doc, "ตารางบันทึกผู้เข้าชมเว็บไซต์ (visitor_logs)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ใช้เก็บสถิติการเข้าชมของผู้ใช้งานเพื่อนำไปวิเคราะห์พฤติกรรมการใช้งานและความนิยมของสถานที่", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: ประกอบด้วยรหัสบันทึก (id: PK), หมายเลขไอพี (ip_address), ข้อมูลเบราว์เซอร์ (user_agent), ประเภทอุปกรณ์ (device_type: Mobile/Desktop), ระบบปฏิบัติการ (os_name), โปรแกรมเปิดเว็บ (browser), หน้าที่เข้าชม (page_path), แหล่งที่มา (referrer), รหัสเซสชัน (session_id) และวันเวลา (timestamp)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: บันทึกหมายเลข 1 มาจากอุปกรณ์ Mobile (iOS - Safari), IP: 127.0.0.1, หน้าเว็บ '/', เวลา 19:40 น.", font_name="Angsana New", size_pt=16, space_after=6)

    doc.add_page_break()

    # หน้า 5: Data Dictionary (ต่อ)
    add_p(doc, "5", font_name="Angsana New", size_pt=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Table 5: user_events
    add_p(doc, "ตารางกิจกรรมการใช้งานบนแผนที่ (user_events)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ทำหน้าที่เก็บประวัติกิจกรรมของผู้ใช้บนแผนผังจำลอง (เช่น ค้นหาอาคาร, กดดูรายละเอียด, กดนำทาง GPS)", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: เชื่อมโยงกับผู้เข้าชมผ่าน visitor_id (FK), ประเภทกิจกรรม (event_type), รายละเอียดกิจกรรม (event_data เช่น ชื่ออาคารที่เลือก) และวันเวลาที่เกิดกิจกรรม (timestamp)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: กิจกรรมประเภท 'navigate' บนอาคาร 'อาคาร 1 เฉลิมพระเกียรติฯ 80 พรรษา' โดยผู้เข้าชมหมายเลข 1", font_name="Angsana New", size_pt=16, space_after=6)

    # Table 6: user_activity_logs
    add_p(doc, "ตารางประวัติกิจกรรมและความปลอดภัย (user_activity_logs)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ใช้บันทึกประวัติการเข้าสู่ระบบและการใช้งานของนักศึกษา บุคลากร และผู้ดูแลระบบเพื่อความปลอดภัย", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: ประกอบด้วยรหัสประจำตัวผู้ใช้ (user_id: FK), ชื่อผู้ใช้งาน (user_name), บทบาทสิทธิ์ (role: student/staff/admin), อีเมล (email), หมายเลขไอพี (ip_address), ข้อมูลอุปกรณ์ (device) และวันเวลา (timestamp)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: แอดมิน 'lnwpoon007x' เข้าสู่ระบบสำเร็จจากอุปกรณ์ Desktop (macOS - Safari) เมื่อเวลา 19:50 น.", font_name="Angsana New", size_pt=16, space_after=6)

    # Table 7: password_resets / password_reset_tokens
    add_p(doc, "ตารางโทเค็นและ OTP กู้คืนรหัสผ่าน (password_resets / password_reset_tokens)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ใช้บันทึกหลักฐานการขอกู้คืนรหัสผ่านและตรวจสอบรหัส OTP 6 หลักผ่านอีเมลมหาวิทยาลัย", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: เก็บประเภทผู้ใช้ (user_type: student/staff), รหัสอ้างอิง (identifier: FK), อีเมลปลายทาง (email), โทเค็นความปลอดภัย (token: UK), รหัสผ่านชั่วคราว (otp), เวลาหมดอายุ (expires_at) และสถานะการใช้งาน (used)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: คำขอกู้รหัสของนักศึกษา 65123456789 รหัส OTP: 849201 หมดอายุภายใน 15 นาที และมีสถานะ Active", font_name="Angsana New", size_pt=16, space_after=6)

    # Table 8: admin_sessions
    add_p(doc, "ตารางเซสชันความปลอดภัยผู้ดูแลระบบ (admin_sessions)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ใช้จัดเก็บโทเค็นการเข้าถึงหน้าจัดการหลังบ้าน (Admin Dashboard) เพื่อป้องกันการเข้าถึงโดยไม่ได้รับอนุญาต", font_name="Angsana New", size_pt=16)
    add_p(doc, "โครงสร้าง: รหัสโทเค็นความปลอดภัย (token: UK), วันเวลาที่สร้าง (created_at), วันเวลาหมดอายุ (expires_at) และสถานะเปิดใช้งาน (is_active)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ตัวอย่างข้อมูล: Admin Token ความยาว 128 ตัวอักษร มีอายุใช้งาน 24 ชั่วโมง และสถานะ Active", font_name="Angsana New", size_pt=16, space_after=8)

    add_p(doc, "   ทุกตารางมีการระบุ Primary Key, Foreign Key และ Unique Key อย่างชัดเจนเพื่อรักษาความสัมพันธ์ของข้อมูล (Data Integrity) และสอดคล้องกับวัตถุประสงค์หลักของระบบแผนผังดิจิทัลและการนำทางอัจฉริยะ มรภ.ศรีสะเกษ", font_name="Angsana New", size_pt=16, bold=True)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 6. หน้า 6+: USER INTERFACE (UI) — Exactly matching 6712732121 format
    # ═════════════════════════════════════════════════════════════════════════
    add_p(doc, "6", font_name="Angsana New", size_pt=14, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, "User Interface (UI)", font_name="Angsana New", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    add_p(doc, "1. หน้าแผนผัง 3D และการสืบค้นข้อมูล (3D Campus Map & Search)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "หน้านี้จะดึงข้อมูลมาจากตาราง buildings เพื่อแสดงผลแผนผังจำลอง 3D ของมหาวิทยาลัยราชภัฏศรีสะเกษในรูปแบบ Interactive ที่เข้าใจง่าย", font_name="Angsana New", size_pt=16)
    add_p(doc, "ลักษณะการออกแบบ: แสดงผลแผนที่ความละเอียดสูง พร้อมหมุด 3D ระบุหมายเลขและชื่ออาคาร มีแผงการ์ด Carousel เลื่อนดูอาคารด้านล่าง", font_name="Angsana New", size_pt=16)
    add_p(doc, "สถานะสีหมวดหมู่อาคาร:", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "• สีน้ำเงิน (Academic): อาคารเรียนและคณะ", font_name="Angsana New", size_pt=16)
    add_p(doc, "• สีส้ม (Office): อาคารสำนักงานและหน่วยงานบริหาร", font_name="Angsana New", size_pt=16)
    add_p(doc, "• สีเขียว (Facility): อาคารศูนย์กีฬาและสิ่งอำนวยความสะดวก", font_name="Angsana New", size_pt=16)
    add_p(doc, "• สีม่วง (Library): อาคารหอสมุดและศูนย์วิทยบริการ", font_name="Angsana New", size_pt=16)
    add_p(doc, "ปุ่มจับพิกัด GPS: เมื่อคลิกปุ่มเป้าเล็ง GPS ด้านขวาบน ระบบจะดึงพิกัดดาวเทียมสดและปักหมุดเรดาร์กระพริบสีฟ้าบนแผนผังทันที", font_name="Angsana New", size_pt=16, space_after=6)

    add_p(doc, "2. หน้ารายละเอียดอาคารและการนำทางอัจฉริยะ (Building Details & Smart Navigation)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "ทำหน้าที่แสดงข้อมูลเชิงลึกของสถานที่ โดยดึงข้อมูลมาจากตาราง buildings", font_name="Angsana New", size_pt=16)
    add_p(doc, "การแสดงผล: แสดงภาพถ่ายอาคาร, ชื่อภาษาไทย/อังกฤษ, หมวดหมู่, เวลาทำการ, เบอร์โทรติดต่อภายใน และพิกัดละติจูดลองจิจูด", font_name="Angsana New", size_pt=16)
    add_p(doc, "ฟังก์ชันการทำงาน:", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "• ปุ่มโทรออก: คลิกเพื่อกดโทรไปยังเบอร์ติดต่อของหน่วยงานภายในอาคารได้ทันที", font_name="Angsana New", size_pt=16)
    add_p(doc, "• ปุ่มนำทาง (Google Maps): เชื่อมโยงพิกัด GPS สดของผู้ใช้ เข้ากับพิกัดอาคารเป้าหมาย และเปิดแท็บ Google Maps นำทางแบบเลี้ยวต่อเลี้ยวทันที", font_name="Angsana New", size_pt=16)
    add_p(doc, "• ปุ่มแชร์พิกัด: คัดลอกลิงก์พิกัดอาคารเพื่อส่งต่อให้เพื่อนหรือผู้มาติดต่อ", font_name="Angsana New", size_pt=16)
    add_p(doc, "การบันทึก: ระบบจะบันทึกกิจกรรมการคลิกและการกดนำทางลงในตาราง user_events โดยอัตโนมัติ", font_name="Angsana New", size_pt=16, space_after=6)

    add_p(doc, "3. หน้าเข้าสู่ระบบและสมัครสมาชิก (Authentication & PDPA Consent)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "หน้านี้รองรับการยืนยันตัวตนของผู้ใช้ทุกระดับ ซึ่งสอดคล้องกับตาราง students, staff_users และ password_reset_tokens", font_name="Angsana New", size_pt=16)
    add_p(doc, "การเข้าสู่ระบบ: แบ่งแท็บชัดเจนระหว่าง 'นักศึกษา' (ใช้รหัสนักศึกษา), 'อาจารย์/บุคลากร' (ใช้อีเมล @sskru.ac.th) และ 'ผู้ดูแลระบบ'", font_name="Angsana New", size_pt=16)
    add_p(doc, "ระบบยินยอม PDPA: ในหน้าสมัครสมาชิก มี Checkbox ให้ผู้ใช้ยินยอมการจัดเก็บข้อมูลส่วนบุคคลด้วยตนเอง (Unchecked by default)", font_name="Angsana New", size_pt=16)
    add_p(doc, "ระบบกู้คืนรหัสผ่าน: มีปุ่ม 'ลืมรหัสผ่าน' สำหรับกรอกอีเมลเพื่อรับรหัส OTP 6 หลัก และตั้งรหัสผ่านใหม่ได้อย่างปลอดภัย", font_name="Angsana New", size_pt=16, space_after=6)

    add_p(doc, "4. หน้าจัดการระบบของผู้ดูแลระบบ (Admin Dashboard & Pin Calibration)", font_name="Angsana New", size_pt=16, bold=True)
    add_p(doc, "แสดงผลลัพธ์การบริหารจัดการฐานข้อมูลทั้งหมดที่ดึงมาจากตาราง buildings, students, staff_users, visitor_logs และ user_events", font_name="Angsana New", size_pt=16)
    add_p(doc, "การจัดการอาคาร: หน้าตาราง CRUD สำหรับเพิ่ม ลบ แก้ไขข้อมูลอาคาร และมีโหมดลากย้ายหมุด (Drag & Drop) ปรับเปลี่ยนพิกัด X, Y บนแผนผังสด", font_name="Angsana New", size_pt=16)
    add_p(doc, "การอนุมัติสิทธิ์: แอดมินสามารถตรวจสอบรายชื่อบุคลากรที่สมัครเข้ามา และกดยืนยันอนุมัติสิทธิ์ (Approve) ได้ทันที", font_name="Angsana New", size_pt=16)
    add_p(doc, "รายงานสถิติ: แสดงกราฟสถิติผู้เข้าชมรายสัปดาห์, สัดส่วนอุปกรณ์ (iOS/Android/Desktop) และจัดอันดับอาคารยอดนิยมที่มีการสืบค้นและนำทางมากที่สุด", font_name="Angsana New", size_pt=16)

    # Save to report file in docs/
    output_path = os.path.join(BASE_DIR, "docs", "SSKRU_Campus_Map_Report.docx")
    doc.save(output_path)
    print(f"Template-matched report generated successfully at: {output_path}")

if __name__ == "__main__":
    create_template_matched_report()
