import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()
    
    # Page Setup (A4 Margins)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Style Defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # ─── COVER PAGE ────────────────────────────────────────────────────────
    p_pre = doc.add_paragraph()
    p_pre.paragraph_format.space_before = Pt(40)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("รายงานการวิเคราะห์และออกแบบระบบ\n(System Analysis and Design Report)")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0D, 0x2C, 0x5E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("โครงการพัฒนาระบบแผนผังดิจิทัลและการนำทางอัจฉริยะ\nมหาวิทยาลัยราชภัฏศรีสะเกษ\n(SSKRU Campus Map & Smart Navigation System)")
    sub_run.font.size = Pt(18)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(0x1A, 0x4F, 0xA0)

    p_mid = doc.add_paragraph()
    p_mid.paragraph_format.space_before = Pt(80)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("คณะ/สาขาวิชา เทคโนโลยีสารสนเทศและวิทยาการคอมพิวเตอร์\nมหาวิทยาลัยราชภัฏศรีสะเกษ")
    info_run.font.size = Pt(16)
    info_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_page_break()

    # ─── CHAPTER 1: บทนำและภาพรวมระบบ ────────────────────────────────────
    h1 = doc.add_heading("บทที่ 1: บทนำและภาพรวมระบบ", level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.runs[0].font.size = Pt(20)
    h1.runs[0].font.bold = True
    h1.runs[0].font.color.rgb = RGBColor(0x0D, 0x2C, 0x5E)

    doc.add_heading("1.1 ความเป็นมาและความสำคัญ", level=2)
    p = doc.add_paragraph(
        "มหาวิทยาลัยราชภัฏศรีสะเกษ มีพื้นที่วิทยาเขตขนาดใหญ่และมีอาคารสถานที่ คณะ หน่วยงาน และสิ่งอำนวยความสะดวกจำนวนมาก "
        "ผู้ใช้งานทั้งนักศึกษาใหม่ บุคลากร และผู้มาติดต่อภายนอก มักประสบปัญหาความยากลำบากในการค้นหาตำแหน่งอาคาร การเดินทาง และการระบุเส้นทาง "
        "จึงได้มีการพัฒนาระบบแผนผังดิจิทัล มรภ.ศรีสะเกษ (SSKRU Campus Map) ขึ้น เพื่อตอบสนองการสืบค้นข้อมูลอาคารสถานที่แบบ 3D และการนำทางผ่านดาวเทียม GPS สด"
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading("1.2 วัตถุประสงค์ของระบบ", level=2)
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.left_indent = Inches(0.25)
    p_obj.add_run("1. เพื่อพัฒนาระบบแผนที่ดิจิทัล 3D และข้อมูลสถานที่ภายในมหาวิทยาลัยราชภัฏศรีสะเกษ\n")
    p_obj.add_run("2. เพื่อพัฒนาระบบระบุพิกัด GPS ตำแหน่งสดของผู้ใช้งาน และเชื่อมต่อระบบนำทาง Turn-by-Turn สู่ Google Maps\n")
    p_obj.add_run("3. เพื่อพัฒนาระบบยืนยันตัวตนและการคุ้มครองข้อมูลส่วนบุคคล (PDPA Consent) สำหรับนักศึกษา บุคลากร และผู้ดูแลระบบ\n")
    p_obj.add_run("4. เพื่อพัฒนาระบบบันทึกประวัติการใช้งานและวิเคราะห์สถิติความนิยมของสถานที่ (Visitor Analytics)")

    doc.add_heading("1.3 ขอบเขตของระบบ", level=2)
    p_scope = doc.add_paragraph()
    p_scope.paragraph_format.left_indent = Inches(0.25)
    p_scope.add_run("• ผู้ใช้งานทั่วไป/นักศึกษา: สามารถค้นหาอาคาร, ดูรายละเอียด, จับพิกัด GPS ปัจจุบัน, ใช้งานระบบนำทาง, และแชร์สถานที่ได้\n")
    p_scope.add_run("• บุคลากร: สามารถลงทะเบียนเข้าใช้งานเพื่อเข้าถึงข้อมูลหน่วยงานและบริการภายใน\n")
    p_scope.add_run("• ผู้ดูแลระบบ (Admin): สามารถจัดการข้อมูลอาคาร (เพิ่ม/ลบ/แก้ไข/ลากย้ายพิกัด), จัดการบัญชีผู้ใช้, และดูสรุปสถิติการใช้งาน")

    doc.add_page_break()

    # ─── CHAPTER 2: การออกแบบฐานข้อมูล ────────────────────────────────────
    h2 = doc.add_heading("บทที่ 2: การวิเคราะห์และออกแบบฐานข้อมูล (Database Design)", level=1)
    h2.paragraph_format.space_before = Pt(12)
    h2.runs[0].font.size = Pt(20)
    h2.runs[0].font.bold = True
    h2.runs[0].font.color.rgb = RGBColor(0x0D, 0x2C, 0x5E)

    doc.add_heading("2.1 ผังแสดงความสัมพันธ์ของข้อมูล (Entity-Relationship Diagram : ER-Diagram)", level=2)
    p_er = doc.add_paragraph(
        "โครงสร้างฐานข้อมูลของระบบประกอบด้วย 8 Entity หลัก โดยมีความสัมพันธ์แบบ 1 ต่อ กลุ่ม (1:N) ระหว่างข้อมูลผู้เยี่ยมชมและประวัติกิจกรรมการใช้งาน ดังแสดงในแผนผัง ER-Diagram ด้านล่างนี้:"
    )
    p_er.paragraph_format.first_line_indent = Inches(0.5)

    # Embed High-Res ER Diagram Image
    img_path = "/Users/monphrakan/Mark_map/images/er_diagram_hd.png"
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(img_path, width=Inches(6.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap_run = p_cap.add_run("รูปที่ 2.1: ผังแสดงความสัมพันธ์ของข้อมูล (Entity-Relationship Diagram)")
        p_cap_run.font.size = Pt(13)
        p_cap_run.font.italic = True
        p_cap_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_heading("2.2 พจนานุกรมข้อมูล (Data Dictionary & Attribute Specifications)", level=2)

    tables_data = [
        ("2.2.1 Entity: Building (ข้อมูลอาคารและสถานที่)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับตาราง (Primary Key)"),
            ("building_id", "Integer", "Unique", "No", "หมายเลขประจำอาคาร (1, 2, 3...)"),
            ("name", "Varchar(255)", "-", "No", "ชื่ออาคารภาษาไทย"),
            ("name_en", "Varchar(255)", "-", "Yes", "ชื่ออาคารภาษาอังกฤษ"),
            ("category", "Varchar(50)", "-", "No", "หมวดหมู่อาคาร (academic, office, facility, library)"),
            ("code", "Varchar(50)", "-", "Yes", "รหัสย่อประจำอาคาร"),
            ("coord_x", "Float", "-", "No", "พิกัดแกน X บนแผนผัง 3D (Leaflet Layer)"),
            ("coord_y", "Float", "-", "No", "พิกัดแกน Y บนแผนผัง 3D (Leaflet Layer)"),
            ("lat", "Float", "-", "Yes", "พิกัดดาวเทียม Latitude GPS จริง"),
            ("lng", "Float", "-", "Yes", "พิกัดดาวเทียม Longitude GPS จริง"),
            ("description", "Text", "-", "Yes", "รายละเอียดและหน่วยงานภายในอาคาร"),
            ("phone", "Varchar(100)", "-", "Yes", "เบอร์โทรศัพท์ติดต่อ"),
            ("tags", "JSON", "-", "Yes", "คำค้นหาที่เกี่ยวข้อง (Search Keywords)"),
            ("image", "Varchar(500)", "-", "Yes", "URL รูปภาพอาคาร"),
            ("created_at", "DateTime", "-", "No", "วันที่และเวลาที่เพิ่มข้อมูล"),
            ("updated_at", "DateTime", "-", "No", "วันที่และเวลาที่แก้ไขข้อมูลล่าสุด")
        ]),
        ("2.2.2 Entity: Student (ข้อมูลบัญชีนักศึกษา)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับตาราง (Primary Key)"),
            ("student_id", "Varchar(20)", "Unique", "No", "รหัสนักศึกษา (เช่น 65123456789)"),
            ("name", "Varchar(255)", "-", "No", "ชื่อ - นามสกุล นักศึกษา"),
            ("year_level", "Integer", "-", "No", "ชั้นปีการศึกษา (1, 2, 3, 4)"),
            ("password_hash", "Varchar(128)", "-", "Yes", "รหัสผ่านที่เข้ารหัสแบบ SHA-256"),
            ("password_plain", "Varchar(128)", "-", "Yes", "รหัสผ่านสำหรับระบบจัดการ"),
            ("is_active", "Boolean", "-", "No", "สถานะการใช้งานบัญชี"),
            ("created_at", "DateTime", "-", "No", "วันที่และเวลาที่ลงทะเบียน")
        ]),
        ("2.2.3 Entity: StaffUser (ข้อมูลบุคลากรและอาจารย์)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับตาราง (Primary Key)"),
            ("username", "Varchar(100)", "Unique", "No", "ชื่อผู้ใช้งานสำหรับเข้าสู่ระบบ"),
            ("email", "Varchar(255)", "Unique", "No", "อีเมลมหาวิทยาลัย (@sskru.ac.th)"),
            ("password_hash", "Varchar(128)", "-", "No", "รหัสผ่านที่เข้ารหัสแบบ SHA-256"),
            ("password_plain", "Varchar(128)", "-", "Yes", "รหัสผ่าน"),
            ("is_active", "Boolean", "-", "No", "สถานะการเปิดใช้งานบัญชี"),
            ("is_approved", "Boolean", "-", "No", "สถานะการอนุมัติสิทธิ์โดยผู้ดูแลระบบ"),
            ("created_at", "DateTime", "-", "No", "วันที่และเวลาที่ลงทะเบียน")
        ]),
        ("2.2.4 Entity: VisitorLog (ข้อมูลผู้เข้าชมเว็บไซต์)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับประวัติ (Primary Key)"),
            ("ip_address", "GenericIP", "-", "Yes", "หมายเลข IP Address ของผู้เข้าชม"),
            ("user_agent", "Text", "-", "Yes", "ข้อมูล Browser และ User-Agent"),
            ("device_type", "Varchar(20)", "-", "No", "ประเภทอุปกรณ์ (mobile, desktop, tablet)"),
            ("os_name", "Varchar(50)", "-", "Yes", "ระบบปฏิบัติการ (iOS, Android, macOS, Windows)"),
            ("browser", "Varchar(50)", "-", "Yes", "เว็บเบราว์เซอร์ที่ใช้งาน"),
            ("page_path", "Varchar(255)", "-", "No", "URL หน้าที่เข้าชม"),
            ("referrer", "Varchar(500)", "-", "Yes", "แหล่งที่มาของลิงก์"),
            ("session_id", "Varchar(64)", "-", "Yes", "รหัสเซสชันประจำตัวผู้เยี่ยมชม"),
            ("timestamp", "DateTime", "-", "No", "วันเวลาที่เข้าชม")
        ]),
        ("2.2.5 Entity: UserEvent (ข้อมูลกิจกรรมการใช้งานบนแผนที่)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับกิจกรรม (Primary Key)"),
            ("visitor_id", "Integer", "FK", "Yes", "อ้างอิงไปยัง VisitorLog.id (Foreign Key)"),
            ("event_type", "Varchar(30)", "-", "No", "ประเภทกิจกรรม (page_view, search, navigate, share)"),
            ("event_data", "Varchar(255)", "-", "Yes", "ข้อมูลของกิจกรรม เช่น ชื่ออาคารที่กดนำทาง"),
            ("timestamp", "DateTime", "-", "No", "วันเวลาที่เกิดกิจกรรม")
        ]),
        ("2.2.6 Entity: UserActivityLog (ประวัติกิจกรรมผู้ใช้งานระบบ)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับประวัติ (Primary Key)"),
            ("user_id", "Varchar(100)", "-", "No", "รหัสประจำตัวผู้ใช้ (รหัสนักศึกษา/Username)"),
            ("user_name", "Varchar(255)", "-", "No", "ชื่อ-นามสกุล ผู้ใช้งาน"),
            ("role", "Varchar(50)", "-", "No", "สิทธิ์การใช้งาน (student, staff, admin)"),
            ("email", "Varchar(255)", "-", "Yes", "อีเมลผู้ใช้งาน"),
            ("ip_address", "Varchar(100)", "-", "Yes", "หมายเลข IP Address"),
            ("device", "Varchar(255)", "-", "Yes", "อุปกรณ์และระบบปฏิบัติการที่ใช้ล็อกอิน"),
            ("timestamp", "DateTime", "-", "No", "วันเวลาที่เกิดกิจกรรม")
        ]),
        ("2.2.7 Entity: PasswordResetToken (โทเค็นและ OTP รีเซ็ตรหัสผ่าน)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับโทเค็น (Primary Key)"),
            ("user_type", "Varchar(20)", "-", "No", "ประเภทผู้ใช้ (student, staff)"),
            ("identifier", "Varchar(100)", "-", "No", "รหัสนักศึกษา หรือ Username"),
            ("email", "Varchar(255)", "-", "No", "อีเมลที่ใช้รับรหัส OTP"),
            ("token", "Varchar(64)", "Unique", "No", "โทเค็นความปลอดภัย"),
            ("otp", "Varchar(10)", "-", "No", "รหัส OTP 6 หลัก"),
            ("expires_at", "DateTime", "-", "No", "เวลาหมดอายุของ OTP"),
            ("used", "Boolean", "-", "No", "สถานะการใช้งานแล้ว"),
            ("created_at", "DateTime", "-", "No", "วันเวลาที่สร้างโทเค็น")
        ]),
        ("2.2.8 Entity: AdminSession (เซสชันผู้ดูแลระบบ)", [
            ("Attribute", "Data Type", "Key", "Null", "คำอธิบาย (Description)"),
            ("id", "Integer (Auto)", "PK", "No", "รหัสลำดับเซสชัน (Primary Key)"),
            ("token", "Varchar(128)", "Unique", "No", "โทเค็นเซสชัน Admin"),
            ("created_at", "DateTime", "-", "No", "วันเวลาที่สร้าง"),
            ("expires_at", "DateTime", "-", "No", "วันเวลาหมดอายุ"),
            ("is_active", "Boolean", "-", "No", "สถานะเซสชันที่เปิดใช้งาน")
        ])
    ]

    for title, rows in tables_data:
        doc.add_heading(title, level=3)
        table = doc.add_table(rows=len(rows), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for idx, text in enumerate(rows[0]):
            hdr_cells[idx].text = text
            set_cell_background(hdr_cells[idx], "1A4FA0")
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(13)
            set_cell_margins(hdr_cells[idx], 80, 80, 100, 100)

        # Data Rows
        for r_idx, row_data in enumerate(rows[1:], start=1):
            row_cells = table.rows[r_idx].cells
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].text = val
                set_cell_background(row_cells[c_idx], bg_color)
                p = row_cells[c_idx].paragraphs[0]
                if c_idx in [0, 2, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(12)
                    if c_idx == 2 and val in ["PK", "FK", "Unique"]:
                        run.font.bold = True
                        if val == "PK":
                            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                        elif val == "FK":
                            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                        else:
                            run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
                set_cell_margins(row_cells[c_idx], 60, 60, 80, 80)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ─── CHAPTER 3: สถาปัตยกรรมและการทำงานของระบบ ─────────────────────────
    h3 = doc.add_heading("บทที่ 3: สถาปัตยกรรมและเทคโนโลยีที่ใช้ในระบบ", level=1)
    h3.paragraph_format.space_before = Pt(12)
    h3.runs[0].font.size = Pt(20)
    h3.runs[0].font.bold = True
    h3.runs[0].font.color.rgb = RGBColor(0x0D, 0x2C, 0x5E)

    doc.add_heading("3.1 สถาปัตยกรรมระบบ (System Architecture)", level=2)
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.left_indent = Inches(0.25)
    p_arch.add_run("1. Frontend Layer: HTML5, Vanilla CSS3 (Design Tokens & Glassmorphism), JavaScript (ES6+), Leaflet.js สำหรับ Interactive Map Engine\n")
    p_arch.add_run("2. Backend Layer: Django Framework (Python 3.13) และ PHP 8.x Architecture API รองรับทั้ง RESTful API และ Server-side Rendering\n")
    p_arch.add_run("3. Database Layer: PostgreSQL / SQLite รองรับทั้งระบบโมเดล ORM และ JSON Fallback Datastores\n")
    p_arch.add_run("4. Deployment & Infrastructure: Vercel Cloud Serverless Architecture รองรับ HTTPS, Proxy SSL, Signed Cookie Sessions และ Service Worker PWA")

    doc.add_heading("3.2 อัลกอริทึมการแปลงพิกัด (Coordinate Calibration Engine)", level=2)
    p_cal = doc.add_paragraph(
        "ระบบใช้สมการถดถอยเชิงเส้น (Linear Interpolation Calibration) ในการแปลงพิกัดดาวเทียม GPS สด (Latitude, Longitude) "
        "ไปเป็นพิกัดพิกเซลบนภาพแผนผัง 3D (Y, X) โดยอ้างอิงจากจุดควบคุมหลักภายในมหาวิทยาลัย:"
    )
    p_cal.paragraph_format.first_line_indent = Inches(0.5)

    p_eq = doc.add_paragraph()
    p_eq.paragraph_format.left_indent = Inches(0.5)
    p_eq.add_run("Y_map = Y1 + KY * (Latitude - LAT1)\n")
    p_eq.add_run("X_map = X1 + KX * (Longitude - LNG1)\n")
    p_eq.runs[0].font.bold = True
    p_eq.runs[0].font.size = Pt(14)
    p_eq.runs[0].font.color.rgb = RGBColor(0x1A, 0x4F, 0xA0)

    doc.add_heading("3.3 การคุ้มครองข้อมูลส่วนบุคคล (PDPA Compliance)", level=2)
    p_pdpa = doc.add_paragraph(
        "ระบบออกแบบให้สอดคล้องกับ พ.ร.บ. คุ้มครองข้อมูลส่วนบุคคล พ.ศ. 2562 (PDPA) โดยมีระบบแจ้งเตือนขอความยินยอม (Consent Modal) "
        "ก่อนการเข้าถึงตำแหน่ง GPS ของอุปกรณ์ และไม่มีการบังคับเลือกในหน้าเข้าสู่ระบบ พร้อมทั้งจัดเก็บประวัติการยินยอมใน Local Storage อย่างปลอดภัย"
    )
    p_pdpa.paragraph_format.first_line_indent = Inches(0.5)

    # Save to project folder
    output_path = "/Users/monphrakan/Mark_map/SSKRU_Campus_Map_Report.docx"
    doc.save(output_path)
    print(f"Report created successfully at: {output_path}")

if __name__ == "__main__":
    create_report()
